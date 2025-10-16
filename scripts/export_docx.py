#!/usr/bin/env python3
"""
Export sections to Word (.docx) document with Chinese translation.

This script reads section files from docs/out/sections and generates
a Word document with English content and placeholder Chinese translation.
It applies glossary terms from config/terms_zh.yaml for consistency.
"""

import sys
import yaml
import re
from pathlib import Path
from typing import Dict, List, Optional
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


def load_glossary(glossary_path: str) -> Dict[str, str]:
    """Load terminology glossary from YAML file."""
    with open(glossary_path, 'r', encoding='utf-8') as f:
        data = yaml.safe_load(f)
    return data.get('terms', {})


def apply_glossary(text: str, glossary: Dict[str, str]) -> str:
    """Apply glossary substitutions to text."""
    result = text
    
    # Sort by length (longest first) to handle multi-word terms properly
    sorted_terms = sorted(glossary.items(), key=lambda x: len(x[0]), reverse=True)
    
    for en_term, zh_term in sorted_terms:
        # Use word boundaries for better matching
        pattern = r'\b' + re.escape(en_term) + r'\b'
        result = re.sub(pattern, zh_term, result, flags=re.IGNORECASE)
    
    return result


def translate_placeholder(text: str, glossary: Dict[str, str]) -> str:
    """
    Generate placeholder Chinese translation.
    
    Currently applies glossary terms to English text as a starting point.
    Future versions will use human translations from docs/translation/sections/*.zh.md
    """
    # For now, apply glossary substitutions to demonstrate the wiring
    translated = apply_glossary(text, glossary)
    return translated


def load_human_translation(section_file: Path, translation_dir: Path) -> Optional[str]:
    """
    Load human translation if available.
    
    Looks for corresponding .zh.md file in docs/translation/sections/
    """
    # Convert 01_Abstract.md -> 01_Abstract.zh.md
    zh_filename = section_file.stem + ".zh.md"
    zh_path = translation_dir / zh_filename
    
    if zh_path.exists():
        content = zh_path.read_text(encoding='utf-8')
        # Remove the heading line if present
        lines = content.split('\n')
        if lines and lines[0].startswith('#'):
            content = '\n'.join(lines[1:]).strip()
        return content
    
    return None


def get_section_files(sections_dir: Path) -> List[Path]:
    """Get all section markdown files sorted by filename."""
    files = sorted(sections_dir.glob("*.md"))
    return files


def create_docx(sections_dir: str, output_path: str, glossary_path: str, translation_dir: str) -> None:
    """Main function to create Word document."""
    sections_dir = Path(sections_dir)
    output_path = Path(output_path)
    glossary_path = Path(glossary_path)
    translation_dir = Path(translation_dir)
    
    print(f"Loading glossary from {glossary_path}...")
    glossary = load_glossary(str(glossary_path))
    print(f"✓ Loaded {len(glossary)} terms")
    
    print(f"Reading section files from {sections_dir}...")
    section_files = get_section_files(sections_dir)
    print(f"✓ Found {len(section_files)} sections")
    
    # Create Word document
    print("Creating Word document...")
    doc = Document()
    
    # Add title
    title = doc.add_heading('Dilithium Digital Signature Scheme', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Add subtitle
    subtitle = doc.add_paragraph('中文翻译版本 / Chinese Translation')
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    subtitle_run = subtitle.runs[0]
    subtitle_run.font.size = Pt(14)
    subtitle_run.font.color.rgb = RGBColor(100, 100, 100)
    
    doc.add_paragraph()  # Blank line
    
    # Process each section
    for section_file in section_files:
        print(f"  Processing {section_file.name}...")
        
        # Read section content
        content = section_file.read_text(encoding='utf-8')
        lines = content.split('\n')
        
        # Extract title (first line starting with #)
        title = None
        content_lines = []
        for line in lines:
            if line.startswith('#') and title is None:
                title = line.lstrip('#').strip()
            else:
                content_lines.append(line)
        
        en_content = '\n'.join(content_lines).strip()
        
        if not title or not en_content:
            continue
        
        # Add section heading
        doc.add_heading(title, 1)
        
        # Check for human translation
        zh_content = load_human_translation(section_file, translation_dir)
        
        if zh_content:
            # Use human translation
            print(f"    ✓ Using human translation")
            
            # Add Chinese translation
            p = doc.add_paragraph(zh_content)
            p.style = 'Normal'
            
            # Add English as reference in smaller font
            doc.add_paragraph()
            en_para = doc.add_paragraph("English Reference:")
            en_run = en_para.runs[0]
            en_run.font.size = Pt(9)
            en_run.font.color.rgb = RGBColor(150, 150, 150)
            en_run.italic = True
            
            en_ref = doc.add_paragraph(en_content)
            for run in en_ref.runs:
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(150, 150, 150)
        else:
            # Use placeholder translation with glossary
            print(f"    → Using placeholder translation")
            zh_content = translate_placeholder(en_content, glossary)
            
            # Add placeholder Chinese (glossary-substituted English)
            p = doc.add_paragraph(zh_content)
            p.style = 'Normal'
            
            # Add note about placeholder
            note = doc.add_paragraph("（待翻译 / Translation pending）")
            note_run = note.runs[0]
            note_run.font.size = Pt(9)
            note_run.font.color.rgb = RGBColor(200, 100, 100)
            note_run.italic = True
        
        doc.add_paragraph()  # Blank line between sections
    
    # Save document
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    print(f"\n✓ Saved Word document to {output_path}")


def main():
    """Main entry point."""
    if len(sys.argv) < 2:
        print("Usage: python export_docx.py [sections_dir] [output_file] [glossary_file] [translation_dir]")
        print("Example: python export_docx.py docs/out/sections docs/out/Dilithium_zh.docx config/terms_zh.yaml docs/translation/sections")
        print("\nDefaults:")
        print("  sections_dir: docs/out/sections")
        print("  output_file: docs/out/Dilithium_zh.docx")
        print("  glossary_file: config/terms_zh.yaml")
        print("  translation_dir: docs/translation/sections")
    
    sections_dir = sys.argv[1] if len(sys.argv) > 1 else "docs/out/sections"
    output_file = sys.argv[2] if len(sys.argv) > 2 else "docs/out/Dilithium_zh.docx"
    glossary_file = sys.argv[3] if len(sys.argv) > 3 else "config/terms_zh.yaml"
    translation_dir = sys.argv[4] if len(sys.argv) > 4 else "docs/translation/sections"
    
    if not Path(sections_dir).exists():
        print(f"Error: Sections directory not found: {sections_dir}")
        print("Run segment_sections.py first to create section files.")
        sys.exit(1)
    
    if not Path(glossary_file).exists():
        print(f"Error: Glossary file not found: {glossary_file}")
        sys.exit(1)
    
    create_docx(sections_dir, output_file, glossary_file, translation_dir)


if __name__ == "__main__":
    main()
